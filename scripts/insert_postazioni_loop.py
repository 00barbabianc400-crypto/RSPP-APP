"""
Inserisce il loop Docxtemplater nella tabella misure postazioni (Table 9, riga dati).
Template in uso in app: Modulo Illuminamento VDT 2022.docx → bucket modelli/MOD_VDT_ILLUMINAMENTO.docx
"""

import os
import sys
from docx import Document

TABLE_INDEX = 9
DATA_ROW_INDEX = 3

# Prima cella: apertura loop + postazione; ultima: annotazioni + chiusura loop
ROW_TAGS = [
    '{{#POSTAZIONI}}{{POSTAZIONE}}',
    '{{N_FINESTRE}}',
    '{{OSCURAMENTO}}',
    '{{LUX_PIANO}}',
    '{{LUX_CENTRO}}',
    '{{ANNOTAZIONI}}{{/POSTAZIONI}}',
]

# Solo template VDT (unico file usato da MOD_VDT_ILLUMINAMENTO in Supabase)
DEFAULT_PATHS = [r'c:\Users\Andrea\Desktop\Modulo Illuminamento VDT 2022.docx']


def set_cell_text(cell, text):
    if not cell.paragraphs:
        cell.text = text
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(text)


def patch_table(doc):
    if len(doc.tables) <= TABLE_INDEX:
        print(f'  SKIP: tabella {TABLE_INDEX} assente ({len(doc.tables)} tabelle)')
        return False
    table = doc.tables[TABLE_INDEX]
    if len(table.rows) <= DATA_ROW_INDEX:
        print(f'  SKIP: riga dati {DATA_ROW_INDEX} assente')
        return False
    row = table.rows[DATA_ROW_INDEX]
    if len(row.cells) < len(ROW_TAGS):
        print(f'  SKIP: attese {len(ROW_TAGS)} colonne, trovate {len(row.cells)}')
        return False
    for i, tag in enumerate(ROW_TAGS):
        set_cell_text(row.cells[i], tag)
    return True


def main():
    paths = sys.argv[1:] if len(sys.argv) > 1 else DEFAULT_PATHS
    if not paths:
        print('Uso: python insert_postazioni_loop.py [file1.docx ...]')
        sys.exit(1)

    for path in paths:
        if not os.path.isfile(path):
            print(f'File non trovato: {path}')
            continue
        print(f'Patch: {path}')
        doc = Document(path)
        if patch_table(doc):
            doc.save(path)
            print('  OK — loop POSTAZIONI inserito nella riga dati')
        else:
            print('  Non modificato')


if __name__ == '__main__':
    main()
